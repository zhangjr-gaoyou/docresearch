"""Document loading and parsing for PDF, DOCX, Markdown."""
from pathlib import Path
from dataclasses import dataclass
from typing import Optional

from pypdf import PdfReader
from docx import Document as DocxDocument


@dataclass
class Document:
    """Unified document structure."""
    content: str
    filename: str
    file_type: str
    metadata: Optional[dict] = None


def load_pdf(file_path: Path) -> Document:
    """Load PDF file and extract text."""
    reader = PdfReader(str(file_path))
    pages = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text)
    content = "\n\n".join(pages)
    return Document(
        content=content,
        filename=file_path.name,
        file_type="pdf",
        metadata={"page_count": len(reader.pages)},
    )


def load_docx(file_path: Path) -> Document:
    """Load DOCX file and extract text."""
    doc = DocxDocument(str(file_path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    content = "\n\n".join(paragraphs)
    return Document(
        content=content,
        filename=file_path.name,
        file_type="docx",
        metadata={},
    )


def load_markdown(file_path: Path) -> Document:
    """Load Markdown file."""
    content = file_path.read_text(encoding="utf-8", errors="replace")
    return Document(
        content=content,
        filename=file_path.name,
        file_type="markdown",
        metadata={},
    )


def load_document(file_path: Path) -> Document:
    """Load document by file extension."""
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        return load_pdf(file_path)
    if suffix in (".docx", ".doc"):
        return load_docx(file_path)
    if suffix in (".md", ".markdown"):
        return load_markdown(file_path)
    raise ValueError(f"Unsupported file type: {suffix}")
